import io
from datetime import datetime, timedelta, timezone
from typing import BinaryIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from apps.soil_laboratory.exceptions import MeasurementsReportGenerationError
from apps.soil_laboratory.models import Measurement
from apps.soil_laboratory.repositories.measurement import MeasurementRepository
from apps.soil_laboratory.schemas.measurement import (
    MeasurementsReportGenerationRequest,
    MeasurementsReportGenerationResponse
)


class MeasurementsReportService:
    """Сервис для работы с отчетами"""

    def __init__(self, measurements_repo: MeasurementRepository):
        self._measurements_repo = measurements_repo

    async def generate_measurements_report(
        self,
        report_request: MeasurementsReportGenerationRequest
    ) -> tuple[MeasurementsReportGenerationResponse, BinaryIO]:
        try:
            # Получаем данные согласно фильтрам
            measurements = await self._get_filtered_measurements(report_request)

            if not measurements:
                raise MeasurementsReportGenerationError("No data to generate measurements report")

            # Генерируем документ
            doc_buffer = await self._create_document(measurements, report_request)

            # Генерируем имя файла
            file_name = self._generate_filename()

            # Создаем ответ
            response = MeasurementsReportGenerationResponse(
                success=True,
                message="Measurements report successfully generated",
                file_name=file_name,
                total_records=len(measurements),
                generated_at=datetime.now(timezone.utc)
            )

            return response, doc_buffer

        except Exception as e:
            raise MeasurementsReportGenerationError(f"Error while generating report: {str(e)}")

    async def _get_filtered_measurements(
        self,
        report_request: MeasurementsReportGenerationRequest
    ) -> list[Measurement]:
        conditions = []

        if report_request.date_from:
            conditions.append(Measurement.created_at >= report_request.date_from)

        if report_request.date_to:
            conditions.append(Measurement.created_at <= report_request.date_to)

        if conditions:
            result = await self._measurements_repo.get_all(where_conditions=conditions)
        else:
            today = datetime.now(timezone.utc).date()
            tomorrow = today + timedelta(days=1)

            conditions = [Measurement.created_at >= today, Measurement.created_at < tomorrow]
            result = await self._measurements_repo.get_all(conditions)

        return result

    async def _create_document(
        self,
        measurements: list[Measurement],
        report_request: MeasurementsReportGenerationRequest
    ) -> BinaryIO:
        """Создает документ Word"""
        doc = Document()

        # Настройка страницы A4
        self._setup_page_format(doc)

        # Добавляем содержимое
        self._add_header(doc, "Звіт")
        self._add_metadata(report_request, doc, len(measurements))
        self._add_table(doc, measurements)
        self._add_signature_section(doc)

        # Сохраняем в буфер
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return doc_buffer

    @staticmethod
    def _setup_page_format(doc: Document) -> None:
        """Настраивает формат страницы A4"""
        sections = doc.sections
        for section in sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)

    @staticmethod
    def _add_header(doc: Document, title: str) -> None:
        """Добавляет заголовок документа"""
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(
            f"Дата формування звіту: {datetime.now().strftime("%d.%m.%Y %H:%M")}"
        )
        date_run.bold = True

        doc.add_paragraph()  # Пустая строка

    @staticmethod
    def _add_metadata(
        report_request: MeasurementsReportGenerationRequest,
        doc: Document,
        measurements_count: int
    ) -> None:
        """Добавляет метаданные документа"""
        count_paragraph = doc.add_paragraph()
        count_run = count_paragraph.add_run(f"Загальна кількість вимірювань: {measurements_count}")
        count_run.bold = True

        period_paragraph = doc.add_paragraph()
        period_date_from_str = (
            report_request.date_from.strftime("%d.%m.%Y")
            if report_request.date_from else None
        )
        period_date_to_str = (
            report_request.date_from.strftime("%d.%m.%Y")
            if report_request.date_from else None
        )

        if period_date_from_str and period_date_to_str:
            period_str = f"{period_date_from_str} - {period_date_to_str}"
        elif period_date_from_str or period_date_to_str:
            period_str = (
                f"з {period_date_from_str}" if period_date_from_str else f"до {period_date_to_str}"
            )
        else:
            period_str = f"за {datetime.now(timezone.utc).date().strftime("%d.%m.%Y")}"

        period_run = period_paragraph.add_run(f"Період: {period_str}")
        period_run.bold = True

        doc.add_paragraph()  # Пустая строка

    @staticmethod
    def _add_table(doc: Document, measurements: list[Measurement]) -> None:
        """Добавляет таблицу с данными"""
        # Определяем количество колонок
        table = doc.add_table(rows=1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Заголовки
        headers = [
            "№",
            "Номер формувального піску",
            "Міцність на стискання, кгс/см²",
            "Газопрникність, од.",
            "Вологість, %",
            "Час",
            "Примітка"
        ]

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

            for paragraph in hdr_cells[i].paragraphs:

                for run in paragraph.runs:
                    run.bold = True

            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Данные
        for idx, measurement in enumerate(measurements, 1):
            row_cells = table.add_row().cells

            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[1].text = measurement.molding_sand_number
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[2].text = f"{measurement.molding_sand_strength_kgf_cm2:.2f}"
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[3].text = f"{measurement.molding_sand_gas_permeability:.0f}"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[4].text = f"{measurement.molding_sand_moisture_percent:.2f}"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[5].text = measurement.created_at.strftime("%d.%m.%y, %H:%M")

            row_cells[6].text = measurement.note if measurement.note else "-"

    @staticmethod
    def _add_signature_section(doc: Document) -> None:
        """Добавляет секцию для подписи"""
        for _ in range(3):
            doc.add_paragraph()

        signature_paragraph = doc.add_paragraph()
        signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_paragraph.add_run("Підпис: _____________________")

    @staticmethod
    def _generate_filename() -> str:
        """Генерирует уникальное имя файла"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"measurements_report_{timestamp}.docx"
